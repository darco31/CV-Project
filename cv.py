from docx import Document
from docx.shared import Inches

document = Document()
# Add profile picture
document.add_picture('profile_pic.jpg', width=Inches(2.0))

# Get input from the user
name = input('What is your name?')
phone_number = input('What is your phone number?')
email = input('What is your email address?')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email)

# About me section
document.add_heading('About Me')

about_me = input("Tell us all about yourself ")
document.add_paragraph(about_me)


# Work experiences
document.add_heading('Work Experience')

p = document.add_paragraph()

company = input('Enter company ')
from_date = input('From Date  ')
to_date = input('To Date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input('Describe your experience at ' + company)

p.add_run(experience_details)

# More experiences
while True:
    more_experience = input(
        'Would you like to add more work experience? Yes or No')

    if more_experience.lower().strip() == 'yes':
        p = document.add_paragraph()

        company = input('Enter company ')
        from_date = input('From Date  ')
        to_date = input('To Date ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input('Describe your experience at ' + company)

        p.add_run(experience_details)

    else:
        break


document.save('cv.docx')
